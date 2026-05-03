"""
build_doc3_fixed.py
Rebuilds 03_UI_Guide_with_Screenshots.docx properly:
- Each section heading is followed immediately by its screenshot (right size, centred)
- Then the explanatory text below the image
- Proper image widths, captions, spacing
- Clean structure throughout
"""

import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR  = r"D:\My-projects\vishal_Sprt_project\Sports_analytics\DOCS"
SHOTS_DIR = os.path.join(DOCS_DIR, "screenshots")
OUT_PATH  = os.path.join(DOCS_DIR, "docx", "03_UI_Guide_with_Screenshots.docx")

# Which screenshot belongs to which section heading (matched by section number)
SECTION_SCREENSHOTS = {
    "1. App Overview":               ("01_app_overview.png",           5.8),
    "2. The Sidebar":                ("02_sidebar_full.png",           2.8),
    "3. Tab 1":                      ("03_tab1_full.png",              5.8),
    "4. Tab 2":                      ("04_tab2_keyword.png",           5.8),
    "5. Tab 3":                      ("05_tab3_augmented.png",         5.8),
    "6. Tab 4":                      ("06_tab4_model_comparison.png",  5.8),
    "7. Tab 5":                      ("07_tab5_leaderboards.png",      5.8),
    "8. Phase 3":                    ("08_tab6_phase3_narrative.png",  5.8),
}

CAPTION_TEXT = {
    "01_app_overview.png":           "Figure 1 — App Overview on first load",
    "02_sidebar_full.png":           "Figure 2 — Sidebar: Phase 1+2+3 subtitle and model controls",
    "03_tab1_full.png":              "Figure 3 — Tab 1: Phase 1 Outcome Prediction",
    "04_tab2_keyword.png":           "Figure 4 — Tab 2: Shot Intelligence (Keyword Rules)",
    "05_tab3_augmented.png":         "Figure 5 — Tab 3: Phase 2 Augmented Outcome",
    "06_tab4_model_comparison.png":  "Figure 6 — Tab 4: Model Comparison",
    "07_tab5_leaderboards.png":      "Figure 7 — Tab 5: Leaderboards & Rankings",
    "08_tab6_phase3_narrative.png":  "Figure 8 — Tab 6: Phase 3 Narrative prototype (2024 final)",
}

# ── helpers ───────────────────────────────────────────────────────────────────

def set_margins(doc):
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.9)
        s.left_margin = s.right_margin = Inches(1.0)


def style_heading(para, level):
    clr = {1: "1F4E79", 2: "2E75B6", 3: "17375E", 4: "404040"}
    sz  = {1: 20,       2: 15,       3: 13,       4: 11}
    for run in para.runs:
        run.font.color.rgb = RGBColor.from_string(clr.get(level, "000000"))
        run.font.size = Pt(sz.get(level, 11))
        run.font.bold = True


def add_image(doc, filename, width_in, caption):
    path = os.path.join(SHOTS_DIR, filename)
    if not os.path.exists(path):
        doc.add_paragraph(f"[Image not found: {filename}]")
        return

    # Image paragraph — centred
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(2)
    run = para.add_run()
    run.add_picture(path, width=Inches(width_in))

    # Caption paragraph
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after  = Pt(10)
    r = cap.add_run(caption)
    r.font.italic = True
    r.font.size   = Pt(9)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Light border line above caption (thin horizontal rule)
    p = cap._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"),  "single")
    top.set(qn("w:sz"),   "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "BBBBBB")
    pBdr.append(top)
    pPr.append(pBdr)


def add_code_para(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.35)
    para.paragraph_format.right_indent = Inches(0.35)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    run = para.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)


def add_table(doc, headers, rows):
    if not headers:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold  = True
                run.font.size  = Pt(8.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd_cell(hdr[i], "1F4E79")

    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        fill  = "DCE6F1" if ri % 2 == 0 else "FFFFFF"
        for i, val in enumerate(row):
            cells[i].text = str(val)[:200]   # trim very long cells
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8.5)
            shd_cell(cells[i], fill)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def shd_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def strip_md(text):
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*",     r"\1", text)
    text = re.sub(r"\*(.*?)\*",          r"\1", text)
    text = re.sub(r"`([^`]+)`",          r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    text = re.sub(r"^#+\s*", "", text)
    return text.strip()


def add_rich_para(doc, raw_text, style=None):
    """Add paragraph with bold/code inline formatting preserved."""
    if style:
        para = doc.add_paragraph(style=style)
    else:
        para = doc.add_paragraph()
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", raw_text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = para.add_run(part[2:-2])
            r.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = para.add_run(part[1:-1])
            r.font.name  = "Courier New"
            r.font.size  = Pt(9)
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            clean = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", part)
            clean = re.sub(r"\*(.*?)\*", r"\1", clean)
            para.add_run(clean)
    return para


def parse_section(doc, lines):
    """
    Parse a block of markdown lines and add to doc.
    Handles: headings, tables, code blocks, bullets, blockquotes, paragraphs.
    Does NOT inject screenshots (handled by caller).
    """
    i = 0
    in_code  = False
    code_buf = []
    in_table = False
    tbl_buf  = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True; code_buf = []
            else:
                in_code = False
                add_code_para(doc, "\n".join(code_buf))
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # Table
        if "|" in line and line.strip().startswith("|"):
            if not in_table:
                in_table = True; tbl_buf = []
            tbl_buf.append(line)
            i += 1; continue
        else:
            if in_table:
                _flush_table(doc, tbl_buf)
                tbl_buf = []; in_table = False

        # HR
        if line.strip() in ("---", "===", "***"):
            i += 1; continue

        # Skip image tags (already handled outside)
        if line.strip().startswith("!["):
            i += 1; continue

        # Skip TOC links and pure anchor links
        if re.match(r"^\d+\. \[", line) or line.strip().startswith("["):
            i += 1; continue

        # Headings inside section (level 3/4)
        if line.startswith("#### "):
            h = doc.add_heading(line[5:].strip(), level=4)
            style_heading(h, 4)
            i += 1; continue
        if line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            style_heading(h, 3)
            i += 1; continue

        # Bullets
        if re.match(r"^[ \t]*[\-\*] ", line):
            text = re.sub(r"^[ \t]*[\-\*] ", "", line)
            add_rich_para(doc, text, style="List Bullet")
            i += 1; continue
        if re.match(r"^[ \t]*\d+\. ", line):
            text = re.sub(r"^[ \t]*\d+\. ", "", line)
            add_rich_para(doc, text, style="List Number")
            i += 1; continue

        # Blockquote
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run('"' + line[2:].strip() + '"')
            r.font.italic = True
            r.font.color.rgb = RGBColor(0x1A, 0x5E, 0x20)
            r.font.size = Pt(10.5)
            i += 1; continue

        # Normal text
        if line.strip():
            add_rich_para(doc, line.strip())

        i += 1

    if in_table and tbl_buf:
        _flush_table(doc, tbl_buf)


def _flush_table(doc, lines):
    data = [l for l in lines if not re.match(r"^\|[\s\|\-:]+\|$", l.strip())]
    if len(data) < 2:
        return
    headers = [c.strip() for c in data[0].strip("|").split("|")]
    rows = []
    for rl in data[1:]:
        cells = [c.strip() for c in rl.strip("|").split("|")]
        while len(cells) < len(headers):
            cells.append("")
        rows.append([strip_md(c) for c in cells[:len(headers)]])
    add_table(doc, [strip_md(h) for h in headers], rows)


# ── Main builder ─────────────────────────────────────────────────────────────

def build():
    md_path = os.path.join(DOCS_DIR, "03_UI_Guide_with_Screenshots.md")
    with open(md_path, encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    set_margins(doc)

    # ── Cover page ────────────────────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Cricket Analytics"); r.font.size = Pt(30)
    r.font.bold = True; r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Document 3 — UI Guide with Screenshots")
    r2.font.size = Pt(17); r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6); r2.font.bold = True

    doc.add_paragraph()
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Complete Walkthrough of the Ultimate Dashboard")
    r3.font.italic = True; r3.font.size = Pt(12)

    doc.add_paragraph()
    p4 = doc.add_paragraph(); p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("App URL:  http://localhost:8501 (or any port — streamlit run app.py)")
    r4.font.name = "Courier New"; r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    doc.add_page_break()

    # ── Screenshot gallery (all screenshots in order, centred, with captions) ─
    h0 = doc.add_heading("Screenshot Gallery — All 8 App Views", level=1)
    style_heading(h0, 1)
    sub = doc.add_paragraph("All screenshots taken live from the running app (1600×950 viewport, wide main layout).")
    sub.paragraph_format.space_after = Pt(14)

    for fname, caption in CAPTION_TEXT.items():
        path = os.path.join(SHOTS_DIR, fname)
        if not os.path.exists(path):
            continue
        # Section label above image
        lbl = doc.add_paragraph()
        lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = lbl.add_run(caption.split("—")[0].strip())
        r.font.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        lbl.paragraph_format.space_before = Pt(14)
        lbl.paragraph_format.space_after  = Pt(4)

        # Image
        ip = doc.add_paragraph(); ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.paragraph_format.space_before = Pt(0)
        ip.paragraph_format.space_after  = Pt(2)
        # sidebar image narrower; others full width
        w = 2.5 if "sidebar" in fname else 6.0
        ip.add_run().add_picture(path, width=Inches(w))

        # Caption
        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(16)
        cr = cp.add_run(caption)
        cr.font.italic = True; cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ── Section-by-section detailed guide ────────────────────────────────────
    # Split markdown into top-level sections (## headings)
    sections = re.split(r"\n(?=## )", content)

    for sec in sections:
        sec_lines = sec.split("\n")
        if not sec_lines:
            continue
        first = sec_lines[0].strip()

        # Top-level title (# )
        if first.startswith("# ") and not first.startswith("## "):
            h = doc.add_heading(first.lstrip("# ").strip(), level=1)
            style_heading(h, 1)
            # Add subtitle / description
            body = [l for l in sec_lines[1:] if l.strip() and not l.startswith("##")]
            for bl in body[:3]:
                if bl.strip() and not bl.startswith("|") and not bl.startswith("#"):
                    parse_section(doc, [bl])
            continue

        # ## Section heading
        if first.startswith("## "):
            title = first[3:].strip()
            # Skip TOC section
            if "Table of Contents" in title:
                continue

            h = doc.add_heading(title, level=2)
            style_heading(h, 2)
            h.paragraph_format.space_before = Pt(14)

            # Find matching screenshot for this section
            shot_file  = None
            shot_width = 5.8
            stem = title.split("—")[0].strip() if "—" in title else title
            for key, (fn, w) in SECTION_SCREENSHOTS.items():
                key_stem = key.split("—")[0].strip()
                if title.startswith(key_stem) or stem.startswith(key_stem):
                    shot_file  = fn
                    shot_width = w
                    break

            # Collect rest of this section's lines (exclude nested ## headings)
            body_lines = sec_lines[1:]

            # Inject screenshot right after heading (before text explanation)
            if shot_file:
                add_image(doc, shot_file, shot_width,
                          CAPTION_TEXT.get(shot_file, shot_file))

            # Now parse all the text / tables / code in this section
            parse_section(doc, body_lines)

            doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Save
    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")
    sz = os.path.getsize(OUT_PATH) // 1024
    print(f"Size: {sz} KB")


if __name__ == "__main__":
    build()
