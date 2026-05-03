"""
build_docx.py
Converts all 4 markdown documentation files to formatted .docx files
with embedded screenshots where relevant.
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS_DIR = r"D:\My-projects\vishal_Sprt_project\Sports_analytics\DOCS"
SHOTS_DIR = os.path.join(DOCS_DIR, "screenshots")
OUT_DIR   = os.path.join(DOCS_DIR, "docx")
os.makedirs(OUT_DIR, exist_ok=True)

MD_FILES = [
    "01_Data_and_Model_Creation.md",
    "02_Phase_Model_Documentation.md",
    "03_UI_Guide_with_Screenshots.md",
    "04_Presentation_Guide.md",
]

# Screenshots to embed in Doc 3 (UI guide)
SCREENSHOT_MAP = {
    "01_app_overview.png":           "Screenshot 1 — App Overview (first load)",
    "02_sidebar_full.png":           "Screenshot 2 — Sidebar: Model Selectors and Sentiment Toggle",
    "03_tab1_full.png":              "Screenshot 3 — Tab 1: Phase 1 Outcome Prediction",
    "04_tab2_keyword.png":           "Screenshot 4 — Tab 2: Shot Intelligence (Keyword Rules)",
    "05_tab3_augmented.png":         "Screenshot 5 — Tab 3: Phase 2 Augmented Outcome",
    "06_tab4_model_comparison.png":  "Screenshot 6 — Tab 4: Model Comparison",
    "07_tab5_leaderboards.png":      "Screenshot 7 — Tab 5: Leaderboards & Rankings",
    "08_tab6_phase3_narrative.png":  "Screenshot 8 — Tab 6: Phase 3 Narrative (2024 final)",
}


def set_heading_style(para, level):
    """Apply heading colour and size."""
    colours = {1: "1F4E79", 2: "2E75B6", 3: "2E75B6", 4: "404040"}
    sizes   = {1: 22, 2: 16, 3: 14, 4: 13}
    for run in para.runs:
        run.font.color.rgb = RGBColor.from_string(colours.get(level, "000000"))
        run.font.size = Pt(sizes.get(level, 11))
        run.font.bold = True


def add_styled_table(doc, headers, rows, style="Table Grid"):
    """Add a nicely styled table."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Set header cell background
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)

    # Data rows
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for para in cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
        # Alternating row shading
        if ridx % 2 == 0:
            for c in cells:
                tc = c._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D6E4F0")
                tcPr.append(shd)

    return table


def add_code_block(doc, text):
    """Add a code/monospace block."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Light grey background simulation via shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F0F0F0")
    pPr.append(shd)


def add_screenshot(doc, filename, caption):
    """Embed a screenshot with caption."""
    path = os.path.join(SHOTS_DIR, filename)
    if not os.path.exists(path):
        doc.add_paragraph(f"[Screenshot not found: {filename}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(path, width=Inches(6.0))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


def parse_and_add_md(doc, md_path, inject_screenshots=False):
    """Parse a markdown file and add content to doc."""
    with open(md_path, encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    i = 0
    in_table = False
    table_lines = []
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ───────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_lines))
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Tables ────────────────────────────────────────────────────────────
        if line.strip().startswith("|") and "|" in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                # process table
                _render_table(doc, table_lines)
                table_lines = []
                in_table = False

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith("# ") and not line.startswith("## "):
            h = doc.add_heading(line[2:].strip(), level=1)
            set_heading_style(h, 1)
            i += 1
            continue
        if line.startswith("## ") and not line.startswith("### "):
            h = doc.add_heading(line[3:].strip(), level=2)
            set_heading_style(h, 2)
            i += 1
            continue
        if line.startswith("### ") and not line.startswith("#### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            set_heading_style(h, 3)
            i += 1
            continue
        if line.startswith("#### "):
            h = doc.add_heading(line[5:].strip(), level=4)
            set_heading_style(h, 4)
            i += 1
            continue

        # ── Horizontal rules ──────────────────────────────────────────────────
        if line.strip() in ("---", "===", "***"):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # ── Screenshot injection (Doc 3 only) ─────────────────────────────────
        if inject_screenshots and line.strip().startswith("!["):
            # Extract filename from markdown image link
            m = re.match(r"!\[.*?\]\(screenshots/(.*?)\)", line.strip())
            if m:
                fname = m.group(1)
                caption = SCREENSHOT_MAP.get(fname, fname)
                add_screenshot(doc, fname, caption)
            i += 1
            continue

        # ── Bullet lists ──────────────────────────────────────────────────────
        if re.match(r"^[\-\*] ", line):
            text = _strip_inline(line[2:].strip())
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, text)
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            text = _strip_inline(re.sub(r"^\d+\. ", "", line).strip())
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, text)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.startswith("> "):
            text = _strip_inline(line[2:].strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run('"' + text + '"')
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x1A, 0x5E, 0x20)
            run.font.size = Pt(11)
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        if line.strip():
            p = doc.add_paragraph()
            _add_formatted_run(p, _strip_inline(line))
        elif not line.strip() and i > 0:
            # empty line → spacing already handled by paragraph spacing
            pass

        i += 1

    # Flush any open table
    if in_table and table_lines:
        _render_table(doc, table_lines)


def _strip_inline(text):
    """Remove most markdown inline syntax, keep the text."""
    # Bold+italic ***text*** or **text**
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*",     r"\1", text)
    text = re.sub(r"\*(.*?)\*",          r"\1", text)
    # Inline code `text`
    text = re.sub(r"`([^`]+)`", r"\1", text)
    # Markdown links [text](url)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    # Headers in text (fallback)
    text = text.lstrip("#").strip()
    return text


def _add_formatted_run(para, text):
    """Add text with bold/italic inline formatting preserved."""
    # Split on **bold** markers
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.font.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            # strip remaining markdown
            clean = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", part)
            clean = re.sub(r"\*(.*?)\*", r"\1", clean)
            para.add_run(clean)


def _render_table(doc, lines):
    """Render a markdown table into a Word table."""
    # Filter separator rows (|---|---|)
    data_rows = [l for l in lines if not re.match(r"^\|[\s\|\-:]+\|$", l.strip())]
    if len(data_rows) < 2:
        return
    headers = [c.strip() for c in data_rows[0].strip("|").split("|")]
    rows = []
    for row_line in data_rows[1:]:
        cells = [c.strip() for c in row_line.strip("|").split("|")]
        # Pad if needed
        while len(cells) < len(headers):
            cells.append("")
        rows.append(cells[:len(headers)])

    # Strip inline markdown from cells
    clean_headers = [_strip_inline(h) for h in headers]
    clean_rows    = [[_strip_inline(c) for c in row] for row in rows]
    add_styled_table(doc, clean_headers, clean_rows)
    doc.add_paragraph()


# ─────────────────────────────────────────────────────────────────────────────
# Per-document builders
# ─────────────────────────────────────────────────────────────────────────────

def build_doc1(out_path):
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    # Cover page
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("Cricket Analytics")
    r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Document 1 — Data Creation & Model Architecture")
    r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    doc.add_paragraph()
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.add_run("ICC T20 World Cup Ball-by-Ball Prediction System").font.italic = True
    doc.add_page_break()

    parse_and_add_md(doc, os.path.join(DOCS_DIR, "01_Data_and_Model_Creation.md"))
    doc.save(out_path)
    print(f"Saved: {out_path}")


def build_doc2(out_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("Cricket Analytics")
    r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Document 2 — Phase-wise Model Documentation")
    r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    doc.add_paragraph()
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.add_run("How Each Phase Works, What It Improves, All Metrics").font.italic = True
    doc.add_page_break()

    parse_and_add_md(doc, os.path.join(DOCS_DIR, "02_Phase_Model_Documentation.md"))
    doc.save(out_path)
    print(f"Saved: {out_path}")


def build_doc3(out_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("Cricket Analytics")
    r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Document 3 — UI Guide with Screenshots")
    r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    doc.add_paragraph()
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.add_run("Complete Walkthrough of the Ultimate Dashboard").font.italic = True
    doc.add_page_break()

    # Embed all screenshots at start with captions
    doc.add_heading("Screenshots — All Tabs", level=1)
    h = doc.paragraphs[-1]
    set_heading_style(h, 1)
    doc.add_paragraph("All screenshots taken live from the running app at http://localhost:8501 (or the port Streamlit prints)")
    doc.add_paragraph()

    for fname, caption in SCREENSHOT_MAP.items():
        doc.add_heading(caption, level=3)
        h = doc.paragraphs[-1]
        set_heading_style(h, 3)
        add_screenshot(doc, fname, caption)

    doc.add_page_break()
    doc.add_heading("Detailed UI Explanation", level=1)
    h = doc.paragraphs[-1]
    set_heading_style(h, 1)
    doc.add_paragraph()

    parse_and_add_md(doc, os.path.join(DOCS_DIR, "03_UI_Guide_with_Screenshots.md"),
                     inject_screenshots=True)
    doc.save(out_path)
    print(f"Saved: {out_path}")


def build_doc4(out_path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("Cricket Analytics")
    r.font.size = Pt(28); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("Document 4 — Presentation Guide")
    r2.font.size = Pt(16); r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    doc.add_paragraph()
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag.add_run("What to Say When You Present — Scripts, Numbers, Q&A").font.italic = True
    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = note.add_run("READ THIS BEFORE YOUR PRESENTATION")
    r3.font.bold = True; r3.font.size = Pt(13)
    r3.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    doc.add_page_break()

    parse_and_add_md(doc, os.path.join(DOCS_DIR, "04_Presentation_Guide.md"))
    doc.save(out_path)
    print(f"Saved: {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_doc1(os.path.join(OUT_DIR, "01_Data_and_Model_Creation.docx"))
    build_doc2(os.path.join(OUT_DIR, "02_Phase_Model_Documentation.docx"))
    build_doc3(os.path.join(OUT_DIR, "03_UI_Guide_with_Screenshots.docx"))
    build_doc4(os.path.join(OUT_DIR, "04_Presentation_Guide.docx"))
    print("\nAll 4 DOCX files created successfully in DOCS/docx/")
